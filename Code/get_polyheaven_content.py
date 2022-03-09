import os
from bs4 import BeautifulSoup
from selenium import webdriver
from time import sleep
from Code.down_vedio import download_videofile
from urllib.request import urlretrieve
from docx import Document
from docx.shared import Inches
from selenium.webdriver import ActionChains
import pandas as pd
# import pyautogui

def get_all_details(out_path:str,mode:str):

        # 获得网页内容
        soup = BeautifulSoup(open("source_html/poly_heaven_HDRIS.html",encoding='utf8'),features="lxml")
        # print(soup)
        # 获得所有a link
        a_link_list = soup.find_all("a", attrs={"class": "GridItem_gridItem__0cuEz"})
        list_title = []
        list_csdn_title = []
        list_desc = []
        list_page_link = []
        list_download_link = []
        list_img_link = []
        for ele in a_link_list:
                cur_resource_title = str(ele["href"]).split("/")[2]
                cur_csdn_title = "%s_游戏开发_VR开发_天空盒子_天空背景_无水印_unitySkybox_高清图片资源_16K_EXR"%cur_resource_title
                cur_desc = "可用于UnityVR开发，3D游戏开发，高清天空盒子Skybox素材，游戏环境背景素材，无水印。使用方法：1-导入Unity后将图片的Shape转换成cube形式，2-创建空Material，并转换成Cube/skybox形式，3-将图片拖入新建的SkyboxMaterial，4-用刚创建的Material代替项目中原本的系统默认Skybox"
                cur_detail_page_link = "https://polyhaven.com"+ele["href"]
                cur_download_link = "https://dl.polyhaven.org/file/ph-assets/HDRIs/exr/16k%2B/"+cur_resource_title+"_16k.exr"
                if len(ele.find_all("img")) > 1:
                        cur_img_link = ele.find_all("img")[1]["src"]
                else:
                        cur_img_link = "No img link"
                list_title.append(cur_resource_title)
                list_csdn_title.append(cur_csdn_title)
                list_desc.append(cur_desc)
                list_page_link.append(cur_detail_page_link)
                list_download_link.append(cur_download_link)
                list_img_link.append(cur_img_link)
        df_result = pd.DataFrame()
        df_result["title"] = list_title
        df_result["csdn_title"] = list_csdn_title
        df_result["desc"] = list_desc
        df_result["page_link"] = list_page_link
        df_result["download_link"] = list_download_link
        df_result["img_link"] = list_img_link
        print(df_result)
        # # 获得MP4视频
        # mp4_source = soup.find(name="video",attrs={"class":"video-player__player"})
        # if mp4_source:
        #     src_url = mp4_source.attrs["src"]
        #     download_videofile(video_links=[src_url],save_path=save_path,names="oculus_video.mp4")
        #
        # # 获得第一幅图
        # right_arrow = driver.find_element_by_xpath("//div[contains(@class,'carousel-arrow--right')]")
        # right_arrow.click()
        # html_pic2 = driver.page_source
        # soup2 = BeautifulSoup(html_pic2, "html.parser")
        # try:
        #     pic2_url = soup2.find(name="div",attrs={"class":"carousel-item"})["style"][23:-3]
        # except:
        #     pic2_url = soup2.find_all(name="div", attrs={"class": "carousel-item"})[1]["style"][23:-3]
        # urlretrieve(pic2_url, os.path.join(save_path, "pic2.jpg"))
        #
        # # 获得第二幅图
        # right_arrow.click()
        # html_pic3 = driver.page_source
        # soup3 = BeautifulSoup(html_pic3, "html.parser")
        # try:
        #     pic3_url = soup3.find(name="div",attrs={"class":"carousel-item"})["style"][23:-3]
        # except:
        #     pic3_url = soup3.find(name="div",attrs={"class":"carousel-item"})[1]["style"][23:-3]
        #
        # urlretrieve(pic3_url,os.path.join(save_path,"pic3.jpg"))
        # # 拿最后一张图片
        # try:
        #     pic4_url = soup3.find(name="div",attrs={"class":"carousel-item"})["style"][23:-3]
        # except:
        #     pic4_url = soup3.find_all(name="div",attrs={"class":"carousel-item"})[1]["style"][23:-3]
        #
        # urlretrieve(pic4_url,os.path.join(save_path,"pic4.jpg"))
        #
        # # 右键选择翻译
        # rightClick = ActionChains(driver)
        # position_element = driver.find_element_by_class_name("app-description__title")
        # rightClick.context_click(position_element).perform()
        # sleep(1)
        # pyautogui.typewrite(['down']*7)
        # pyautogui.typewrite(["enter"])
        # sleep(3)
        #
        # # 滚动条加载到最底部
        # total_height = driver.execute_script("return document.body.scrollHeight")
        # new_height = 0
        # print("total_height",total_height)
        # while new_height < total_height:
        #     new_height += 100
        #     driver.execute_script("window.scrollTo(0, %s);"%new_height)
        # # 重新获得翻译后网页内容
        # html = driver.page_source
        # # 获得需要html后开始用soup获取资源
        # soup = BeautifulSoup(html, "html.parser")
        #
        # # 获得题目
        # cur_title_ele = soup.find(name="div",attrs={"class":"app-description__title"})
        # cur_title = cur_title_ele.text
        # print(cur_title)
        #
        # cur_title_ele = soup.find(name="div", attrs={"class": "app-description__title"})
        # cur_cn_title = cur_title_ele.text
        # # 获得内容
        # cur_content = soup.find(name="div",attrs={"class":"store-item-detail-page-description__content"}).text
        # print(cur_content)
        # # 获得评论题目
        # cur_evaluate = soup.find(name="h1",attrs={"class":"bxHeading bxHeading--level-5 app-review__title"}).text
        # print(cur_evaluate)
        # cur_evaluate_content = soup.find(name="p",attrs={"class":"bxParagraph bxParagraph--level-1 app-review__description"}).text
        # print(cur_evaluate_content)
        # # 获得参数信息
        # cur_detail_left = [ele.text for ele in soup.find_all(name="div",attrs={"class":"app-details-row__left"})]
        # cur_detail_right = [ele.text for ele in soup.find_all(name="div",attrs={"class":"app-details-row__right"})]
        # print(cur_detail_left,cur_detail_right)
        #
        # # 形成文章
        # opening ="大家好，我是痴迷VR世界，吹爆机皇Quest的幻境流浪汉白马，有好玩的，不分享怎么行，为你献上好玩的有用的VR游戏情报。今天给大家带来的游戏是-%s(%s)"%(cur_title,org_title)
        # ending = "本期游戏介绍完毕，白马要去发现新的游玩项目了，这次的游戏%s(%s)，你觉得怎么样，如果玩过后有心得体验或者操作秘技，甚至是攻略提问，欢迎在评论中交流。有机会的话希望和你在虚拟世界中相遇。玩儿去喽:)"%(cur_title,org_title)
        #
        # document = Document()
        # document.add_heading("%s(%s)"%(cur_title,org_title),0)
        # document.add_picture("first.png", width=Inches(5))
        # document.add_paragraph(opening)
        #
        # document.add_heading("这是个什么游戏？",1)
        # document.add_picture(os.path.join(save_path,"pic2.jpg"),width=Inches(5))
        # document.add_paragraph(cur_content)
        #
        # document.add_heading("游戏参数",1)
        # document.add_picture(os.path.join(save_path,"pic3.jpg"),width=Inches(5))
        # document.add_paragraph(cur_detail_left)
        # document.add_paragraph(cur_detail_right)
        #
        # document.add_heading("玩后体验",1)
        # document.add_picture(os.path.join(save_path,"pic4.jpg"),width=Inches(5))
        # document.add_paragraph(cur_evaluate)
        # document.add_paragraph(cur_evaluate_content)
        #
        # document.add_paragraph(ending)
        #
        # document.save(os.path.join(save_path,"%s.docx"%os.path.split(save_path)[-1]))

        # driver.close()
# get_all_detail_page("http://oculus.com//experiences/quest/2299465166734471?ranking_trace=106637288010653_2299465166734471_SKYLINEWEBQUESTSTORE_1888816384764129%3D%3D2x5u9EGNTYgMQt5Wi",r"D:\Document\生产\VR\auto_video\\")